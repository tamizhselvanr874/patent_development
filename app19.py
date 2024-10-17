import fitz  # PyMuPDF for PDF extraction  
from openai import AzureOpenAI  
from dotenv import load_dotenv  
import os  
import re  # For parsing structured output  z
import json  # For JSON handling  
import pandas as pd  
import streamlit as st  
import docx  
from io import BytesIO  
from azure.ai.formrecognizer import DocumentAnalysisClient  
from azure.core.credentials import AzureKeyCredential  
from azure.core.exceptions import HttpResponseError  
from docx2pdf import convert  
import pypandoc  
from PyPDF2 import PdfMerger  
import os  
import tempfile 

# Load environment variables from .env file  
load_dotenv()  
# Initialize global variables  
domain_subject_matter = "default domain"  
experience_expertise_qualifications = "default qualifications"  
style_tone_voice = "default style"
  
# Set up Azure OpenAI API credentials from .env  
client = AzureOpenAI(  
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),  # Pull from environment  
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),  # Pull from environment  
    api_version=os.getenv("OPENAI_API_VERSION"),  # Pull from environment  
)  
  
# Azure Form Recognizer setup  
form_recognizer_endpoint = os.getenv("FORM_RECOGNIZER_ENDPOINT")  
form_recognizer_api_key = os.getenv("FORM_RECOGNIZER_API_KEY")  
  
  
def extract_text_from_docx(uploaded_docx):  
    """Extract text from a DOCX file."""  
    doc = docx.Document(uploaded_docx)  
    full_text = []  
    for para in doc.paragraphs:  
        full_text.append(para.text)  
    return "\n".join(full_text)  
  
def determine_domain_expertise(action_document_text):  
    """Analyze the action document to determine the required domain expertise, experience, and analysis style."""  
    global domain_subject_matter, experience_expertise_qualifications, style_tone_voice
    prompt = f"""  
    Analyze the following action document text and determine the domain expertise required to analyze this document:  
    {action_document_text}  
  
    Step 1: Identify the subject matter and domain expertise required to understand this document and the cited documents in depth.  
    Step 2: Determine the experience, expertise, and educational qualifications required to handle this document and the cited documents in depth.  
    Step 3: Describe the style, tone, and voice required to analyze these kinds of documents.  
    NOTE: Each answer needs to be detailed.  
    Step 4: Provide the response in the following JSON format:  
    {{  
        "domain_subject_matter": " Detailed description of the domain subject matter",  
        "experience_expertise_qualifications": "Detailed description of the experience, expertise, and educational qualifications required",  
        "style_tone_voice": "Detailed description of the style, tone, and voice required"  
    }}  
    """  
  
    messages = [  
        {  
            "role": "system",  
            "content": "You are an AI assistant that can analyze the following action document text and determine the domain, expertise, and subject matter required to analyze this document."  
        },  
        {  
            "role": "user",  
            "content": prompt  
        }  
    ]  
  
    # Call OpenAI API for domain expertise determination  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
  
        # Extract the content from the response  
        content = response.choices[0].message.content.strip()  
  
        if content.startswith("```json"):  
            content = content[7:-3].strip()  
        elif content.startswith("```"):  
            content = content[3:-3].strip()  
  
        # Print the raw response for debugging  
        st.write("Raw API Response:")  
        st.write(response.choices[0].message.content)  
  
        # Clean the content to remove invalid control characters  
        cleaned_content = ''.join(c for c in content if c.isprintable())  
  
        # Print the cleaned content for debugging  
        st.write("Cleaned API Response:")  
        st.write(cleaned_content)  
  
        # Attempt to parse the cleaned JSON content  
        try:  
            data = json.loads(cleaned_content)  
        except json.JSONDecodeError as json_err:  
            # If JSON parsing fails, print the error and the cleaned content for debugging  
            st.error(f"JSON decoding error: {json_err}")  
            st.error(f"Cleaned response content: {cleaned_content}")  
            return (None, None, None)  
  
        domain_subject_matter = data.get("domain_subject_matter")  
        experience_expertise_qualifications = data.get("experience_expertise_qualifications")  
        style_tone_voice = data.get("style_tone_voice")  
  
        # Return the results as a tuple  
        return (domain_subject_matter, experience_expertise_qualifications, style_tone_voice)  
    except Exception as e:  
        st.error(f"Error during domain expertise determination: {e}")  
        return (None, None, None)  
  
def check_for_conflicts(action_document_text, domain, expertise, style):
    """
    Analyzes the action document and extracts:
    - Foundational claim
    - Referenced documents
    - Figures and technical text related to them
    """
    global domain_subject_matter, experience_expertise_qualifications, style_tone_voice
    
    # Escape curly braces in the action_document_text
    escaped_text = action_document_text.replace("{", "{{").replace("}", "}}")

    # The content with placeholders dynamically filled
    content = f"""
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """
    # Print the content to the terminal
    print("Generated Content for LLM:\n")
    print(content)

    # Formulate the prompt to be sent to the LLM
    prompt = f"""
    Analyze the following action document text and extract the foundational claim:
    {escaped_text}
    Step 1: Extract the key claims from the document and name it as 'Key_claims'.
    Step 2: From the 'Key_claims' extract the foundational claim and store it in a variable called "foundational_claim" (Note: method claims and system claims are not considered independent claims and only one claim can be the foundational claim).
    Step 3: From the foundational claim, extract the information under U.S.C 102 and/or 103.
    Step 4: Extract all referenced documents under U.S.C. 102 and/or 103 mentioned in the action document specified only in the "foundational_claim".
    Step 5: For each referenced document, create a variable that stores the document name.
    Step 6: If the foundational claim refers to the referenced documents, extract the entire technical content with its specified paragraph location and image reference. Map the claim with the conflicting document name.
    Step 7: Do not extract any referenced document data that is not related to the foundational claim.   
    NOTE: Extract in English.
    Step 8: Return the output as a JSON object with the following structure:
    {{
        "foundational_claim": "text",
        "documents_referenced": ["doc1", "doc2", ...],
        "figures": ["fig1", "fig2", ...],
        "text": "detailed text"
    }}
    """

    messages = [
        {
            "role": "system",
            "content": content  # dynamically generated content for the LLM's role
        },
        {
            "role": "user",
            "content": prompt,  # prompt asking the user for conflict analysis
        },
    ]

    # Call the OpenAI API for conflict checking (assuming you have client setup)
    try:
        response = client.chat.completions.create(
            model="GPT-4-Omni", messages=messages, temperature=0.2
        )
        # Extract the content and remove the triple backticks if necessary
        content = response.choices[0].message.content.strip()

        if content.startswith("```json"):
            content = content[7:-3].strip()
        elif content.startswith("```"):
            content = content[3:-3].strip()

        # Print the raw response for debugging
        print(f"Raw response: {response.choices[0].message.content}")

        # Parse the JSON content
        return json.loads(content)

    except json.JSONDecodeError as e:
        print(f"JSON decoding error: {e}")
        print(f"Raw response: {response.choices[0].message.content}")
        return None
    except Exception as e:
        print(f"Error during conflict checking: {e}")
        return None


# Function to extract and analyze figure-related details  
def extract_figures_and_text(conflict_results, ref_documents_texts, domain, expertise, style):  
    """  
    Extract figures and related technical text from the 'check_for_conflicts' function's output.  
    """  
    # Extract the 'figures' and 'text' sections from the JSON output  
    fig_details = conflict_results.get("figures", [])  
    text_details = conflict_results.get("text", "")  
    content = f"""
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """
    # Print the content to the terminal
    print("Generated Content for LLM:\n")
    print(content)

    # Prepare a structured prompt for figure analysis  
    figure_analysis_prompt = f"""  
    Analyze the figures and technical text from the referenced document in relation to the foundational claim.  
    Instructions:  
    1. Identify Figures:  
        - For each figure referenced in the foundational claim, extract the following:  
            - **Figure Number and Title:** Provide the figure number and its title.  
            - **Technical Details:** Extract all technical details related to the figure as mentioned in the text. Ensure no technical detail is missed.  
            - **Importance:** Explain the importance of the figure in relation to the foundational claim. Describe how it supports, illustrates, or contradicts the claim.  
    2. Extract Text from Paragraphs:  
        - From the paragraphs cited in the foundational claim, extract the relevant text as in the document uploaded and store it in a separate variable.  
    3. Workflow for Cases with Images:  
        - If figures are present in the referenced document:  
            - Follow the steps outlined above to extract figure details and technical information.  
            - Ensure that any interpretations of the figures include specific references to the data or concepts depicted.  
    4. Workflow for Cases without Images:  
        - If no figures are present:  
            - Focus on extracting and analyzing the text from the referenced document.  
            - Identify and highlight key technical details and concepts that are essential to understanding the foundational claim.  
    Input Details:  
    Figures: {json.dumps(fig_details, indent=2)}  
    Text: {text_details}  
    Referenced Document Texts: {json.dumps(ref_documents_texts, indent=2)}  
    Example Output Format:  
    {{  
        "figures_analysis": [  
            {{  
                "figure_number": "Figure 1",  
                "title": "Title of Figure 1",  
                "technical_details": "Detailed text",  
                "importance": "Explanation of importance"  
            }},  
            ...  
        ],  
        "extracted_paragraphs": [  
            "Paragraph text 1",  
            ...  
        ]  
    }}  
    """  
  
    messages = [  
        {  
            "role": "system",  
            "content": content,
        },  
        {  
            "role": "user",  
            "content": figure_analysis_prompt,  
        },  
    ]  
  
    # Call OpenAI API for figure analysis  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
        analysis_output = response.choices[0].message.content.strip()  
  
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
  
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
  
        # Parse the JSON content  
        return json.loads(analysis_output)  
  
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error during figure analysis: {e}")  
        return None  


def extract_details_from_filed_application(filed_application_text, foundational_claim, domain, expertise, style):  
    """  
    Extract details from the filed application related to the foundational claim.  
    """
    content = f"""
   You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """  
    prompt = f"""  
    Analyze the following filed application text and extract details related to the foundational claim.  
    Filed Application Text: {filed_application_text}  
    Foundational Claim: {json.dumps(foundational_claim, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the filed application that relate to the foundational claim.  
    2. Ensure that any extracted details include specific references to the paragraphs or sections in the filed application where they are found. NOTE: Extract in English.  
    3. Return the extracted details in the following JSON format:  
    {{  
        "foundational_claim_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Detailed text related to the foundational claim"  
            }},  
            ...  
        ]  
    }}  
    """  
  
    messages = [  
        {  
            "role": "system",  
            "content": content,
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
  
    # Call OpenAI API for extracting details from the filed application  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
        analysis_output = response.choices[0].message.content.strip()  
  
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
  
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
  
        # Parse the JSON content  
        return json.loads(analysis_output)  
  
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error extracting details from filed application: {e}")  
        return None  

  
# Function to extract details from pending claims and modify the filed application details  
def extract_and_modify_filed_application(filed_application_details, pending_claims_text, domain, expertise, style):  
    """  
    Extract details from the pending claims and modify the filed application details.  
    """
    content = f"""
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """  
    
    global domain_subject_matter, experience_expertise_qualifications, style_tone_voice   
    prompt = f"""  
    Analyze the following pending claims text and modify the filed application details accordingly.  
    Pending Claims Text: {pending_claims_text}  
    Filed Application Details: {json.dumps(filed_application_details, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the pending claims that relate to the foundational claim.  
    2. Modify the filed application details based on the extracted details from the pending claims.  
    3. Ensure that any modifications include specific references to the paragraphs or sections in the pending claims where they are found.NOTE:Extract in English.  
    4. Return the modified filed application details in the following JSON format:  
    {{  
        "modified_filed_application_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Modified detailed text based on pending claims"  
            }},  
            ...  
        ]  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": content,
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for extracting and modifying filed application details  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
          
        # Parse the JSON content  
        return json.loads(analysis_output)  
      
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error extracting and modifying filed application details: {e}")  
        return None  
  
# Function to analyze the filed application based on the foundational claim, figure analysis, and application details  
def analyze_filed_application(extracted_details, foundational_claim, figure_analysis, domain, expertise, style): 
    content = f"""
   You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """ 
    prompt = f"""  
    Analyze the filed application based on the foundational claim:  
    {json.dumps(foundational_claim, indent=2)}  
    and the figure analysis results:  
    {json.dumps(figure_analysis, indent=2)}  
    and the application as filed details:  
    {extracted_details}  
    Assess whether the examiner's rejection of the application under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified by comparing it with the cited references text.  
    Instructions:  
    Key Features of Foundational Claim:

Extract and list the key features of the foundational claim.
Ensure to include structural details, functional aspects, and any specific configurations mentioned in the claim.

Key Features of Cited Reference:
Extract and list the key features of the cited reference.(also include where it is located in the cited text)
Highlight any similarities or differences in structure, function, and configuration compared to the foundational claim.

Examiner’s Analysis:
Describe the examiner’s analysis and the basis for rejection.
Summarize how the examiner interprets the cited reference in relation to the foundational claim.
Identify whether the rejection is based on U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness).

Novelty Analysis (U.S.C 102 - Lack of Novelty):
Compare the foundational claim with the cited reference to determine if the claim lacks novelty.
Identify if all elements of the foundational claim are disclosed in the cited reference.
Provide a detailed side-by-side comparison of each element.

Non-Obviousness Analysis (U.S.C 103 - Obviousness):
Analyze whether the foundational claim is obvious in light of the cited reference.
Consider if the combination of features in the foundational claim would have been obvious to a person skilled in the art at the time of the invention.
Discuss any differences that might contribute to non-obviousness.

Conclusion:
Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified.
Summarize the key points that support or refute the examiner’s rejection. 
    Potential Areas for Distinction Listed  
    Identify potential areas for distinction in the foundational claim. 
    INSTRUCTIONS TO BE FOLLOWED WHILE PROPOSING AMENDMENTS:
Guidance for Proposing Amendments and Arguments:
When proposing amendments:

Be Specific: Clearly identify which feature you are amending and provide detailed enhancements.
Highlight Novel Elements: Emphasize new details such as specific materials, unique configurations, or innovative steps that are not present in the cited reference.
Refer to Sources: Cite sections of the application or figures from which the amendments and supporting arguments are drawn to reinforce their basis.
Maintain Claim Integrity: Ensure that the proposed amendments do not alter the fundamental essence of the original claim but enhance its patentability.
When crafting arguments to the examiner:
Address Rejection Points: Directly counter the examiner's reasons for rejection by highlighting differences between the amended claim and the cited reference.
Emphasize Novelty and Non-Obviousness: Explain why the amended features are new and not obvious, providing clear distinctions from the prior art.
Use Supporting Evidence: Reference specific examples, embodiments, or descriptions in the application that support your arguments.
Be Persuasive: Articulate the advantages and unique aspects of the invention that merit patent protection.
    Proposed Amendments and Arguments  
    Amendment to Foundational Claim for Each Key Feature Listed Separately with the New Features Highlighted:  
        Amendment [Number]: [Feature]  
        Original: "[Original feature description...]"  
        Proposed: "[Enhanced feature description with new details, specific materials, or configurations...]"  

 
        Provide arguments for novelty and non-obviousness over the cited reference.  
    Identify Limitations in the Current Claims  
    Identify limitations in the current claims and propose specific language or structural changes that address those limitations.  
    Propose New Arguments or Amendments  
    Suggest additional arguments or amendments that further distinguish the foundational claim from the prior art and strengthen the application. Include multiple amendments for thorough differentiation.
    FORMATTING NOTES:
     - Ensure the amendments maintain the original intent of the claims while improving clarity and scope.  
      - Do the numbering in bullets and not in numbers. Do not use markdown formatting in your response.  
      -Wherever U.S.C 102 is mentioned, it should be printed as U.S.C 102 (Lack of Novelty), and wherever U.S.C 103 is mentioned, it should be printed as U.S.C 103 (Obviousness).  
     - Bold the key points.
    """  
  
    messages = [  
        {  
            "role": "system",  
            "content": content,
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
        analysis_output = response.choices[0].message.content.strip()  
  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
  
        try:  
            return json.loads(analysis_output)  
        except json.JSONDecodeError:  
            return analysis_output  
    except Exception as e:  
        print(f"Error during filed application analysis: {e}")  
        return None  
 
  
def analyze_modified_application(cited_references_text, foundational_claim, figure_analysis, modified_application_details, domain, expertise, style): 
    content = f"""
   You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """ 
    prompt = f"""  
    Analyze the modified application based on the foundational claim:{json.dumps(foundational_claim, indent=2)}and the figure analysis results:{json.dumps(figure_analysis, indent=2)}and the modified application details:{json.dumps(modified_application_details, indent=2)}and the cited references:{json.dumps(cited_references_text, indent=2)}  
Assess whether the examiner's rejection of the application under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified by comparing it with the cited references text.

Key Features of Foundational Claim:

Extract and list the key features of the foundational claim.
Ensure to include structural details, functional aspects, and any specific configurations mentioned in the claim.

Key Features of Cited Reference:
Extract and list the key features of the cited reference.(also include where it is located in the cited text)
Highlight any similarities or differences in structure, function, and configuration compared to the foundational claim.

Examiner’s Analysis:
Describe the examiner’s analysis and the basis for rejection.
Summarize how the examiner interprets the cited reference in relation to the foundational claim.
Identify whether the rejection is based on U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness).

Novelty Analysis (U.S.C 102 - Lack of Novelty):
Compare the foundational claim with the cited reference to determine if the claim lacks novelty.
Identify if all elements of the foundational claim are disclosed in the cited reference.
Provide a detailed side-by-side comparison of each element.

Non-Obviousness Analysis (U.S.C 103 - Obviousness):
Analyze whether the foundational claim is obvious in light of the cited reference.
Consider if the combination of features in the foundational claim would have been obvious to a person skilled in the art at the time of the invention.
Discuss any differences that might contribute to non-obviousness.

Conclusion:
Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified.
Summarize the key points that support or refute the examiner’s rejection.

Potential Areas for Distinction:
Identify areas where the foundational claim can be distinguished from the cited reference.
Focus on unique structural features, specific materials, configurations, or functions not disclosed in the cited reference.

Proposed Amendments and Arguments:
For each key feature point in the foundational claim, propose specific amendments separately. NOTE: for all the points in the foundational claim, it is mandatory to propose amendments.
Present original and proposed versions, highlighting new features, specific materials, or configurations.
Amendment [Number]: [Feature]
Original: "[Original feature description...]"
Proposed: "[Enhanced feature description with new details, specific materials, or configurations...]"
Provide arguments supporting novelty and non-obviousness over the cited reference.
Emphasize any technical advantages or improvements introduced by the amendments.
NOTE WHILE PROPOSING ARGUMENTS:
'''\Guidance for Proposing Amendments and Arguments:
When proposing amendments:

Be Specific: Clearly identify which feature you are amending and provide detailed enhancements.
Highlight Novel Elements: Emphasize new details such as specific materials, unique configurations, or innovative steps that are not present in the cited reference.
Refer to Sources: Cite sections of the application or figures from which the amendments and supporting arguments are drawn to reinforce their basis.
Maintain Claim Integrity: Ensure that the proposed amendments do not alter the fundamental essence of the original claim but enhance its patentability.
When crafting arguments to the examiner:
Address Rejection Points: Directly counter the examiner's reasons for rejection by highlighting differences between the amended claim and the cited reference.
Emphasize Novelty and Non-Obviousness: Explain why the amended features are new and not obvious, providing clear distinctions from the prior art.
Use Supporting Evidence: Reference specific examples, embodiments, or descriptions in the application that support your arguments.
Be Persuasive: Articulate the advantages and unique aspects of the invention that merit patent protection.\'''

Identify Limitations in Current Claims:
Identify any limitations or weaknesses in the current claims.
Propose specific language or structural changes to address these limitations.
Ensure that the proposed changes do not alter the original intent of the claims.

Propose New Arguments or Amendments:
Suggest additional arguments or amendments to further distinguish the foundational claim from the cited prior art.
Include multiple amendments for thorough differentiation.
Ensure that the original intent of the claims is maintained while improving clarity and scope.
NOTE:
Numbering and Formatting:
Use bullet points (•) instead of numbers when listing items.
Do not include markdown formatting in your response.
Bolden the key points.
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": content,

        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        try:  
            return json.loads(analysis_output)  
        except json.JSONDecodeError:  
            return analysis_output  
    except Exception as e:  
        print(f"Error during modified application analysis: {e}")  
        return None  
  
def save_analysis_to_word(analysis_output):
    if analysis_output is None or analysis_output.strip() == "":
        st.error("Analysis data is missing or empty.")
        return None

    # Create a new Word document
    doc = docx.Document()
    doc.add_heading('Filed Application Analysis Results', level=1)

    # Split the analysis output into lines
    lines = analysis_output.split('\n')
    for line in lines:
        line = line.strip()

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
        else:
            # Create a new paragraph for normal or mixed text (bold and non-bold)
            paragraph = doc.add_paragraph()

            # Use regex to find text between **...** for bold words
            # Split by bold sections while keeping bold markers for processing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    # This is the bold part, remove the '**' and set it as bold
                    bold_text = part[2:-2]
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                else:
                    # This is regular text
                    paragraph.add_run(part)

    # Save the document to a BytesIO buffer instead of writing to disk
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Define your endpoint and key  
form_recognizer_endpoint = "https://patentocr.cognitiveservices.azure.com/"  # Replace with your actual endpoint  
form_recognizer_api_key = "cd6b8996d93447be88d995729c924bcb"
# Initialize session state variables  
session_vars = [  
    'conflict_results', 'foundational_claim', 'figure_analysis', 'filed_application_analysis',  
    'cited_documents', 'pending_claims_analysis', 'pending_claims_available', 'domain', 'expertise',  
    'style', 'filed_application_name'  
]  
  
for var in session_vars:  
    if var not in st.session_state:  
        st.session_state[var] = None  
  
st.session_state['pending_claims_available'] = st.session_state.get('pending_claims_available', "No")  

# Function to create aligned uploader and button  
def create_uploader_and_button(label_button, key):  
    col1, col2 = st.columns([4, 1])  # Adjust the column widths as needed  
    with col1:  
        uploaded_file = st.file_uploader("", type=["pdf", "docx"], key=key)  # Empty string for no label  
    with col2:  
        st.markdown("<br>", unsafe_allow_html=True)  # Add some space with HTML  
        button_clicked = st.button(label_button)  
    return uploaded_file, button_clicked  
  
def convert_docx_to_pdf(docx_path, pdf_path):  
    """Convert a DOCX file to PDF using docx2pdf."""  
    try:  
        convert(docx_path, pdf_path)  
        return pdf_path  
    except Exception as e:  
        st.error(f"Failed to convert DOCX to PDF: {e}")  
        return None  
  
def extract_text_from_pdf(uploaded_pdf_path):  
    """Extract text from a PDF file using Azure Form Recognizer Document Intelligence."""  
    try:  
        # Initialize DocumentAnalysisClient  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
  
        # Read the file content  
        with open(uploaded_pdf_path, "rb") as f:  
            file_content = f.read()  
  
        # Use the prebuilt-document model to analyze the document  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
  
        # Get the result of the analysis  
        result = poller.result()  
  
        # Extract the text from the result  
        text = ""  
        for page in result.pages:  
            for line in page.lines:  
                text += line.content + "\n"  
  
        return text  
  
    except HttpResponseError as e:  
        st.error(f"Failed to analyze the document: {e.message}")  
        return None  
  
    except Exception as e:  
        st.error(f"An unexpected error occurred: {e}")  
        return None  
  
# Function to convert DOCX to PDF  
def convert_word_to_pdf(input_file, output_file):  
    try:  
        pypandoc.convert_file(input_file, 'pdf', outputfile=output_file, extra_args=['--pdf-engine=pdflatex'])  
        return output_file  
    except Exception as e:  
        st.error(f"Error converting file: {e}")  
        return None  
  
# Function to merge multiple PDFs  
def merge_pdfs(pdf_list, output_file):  
    merger = PdfMerger()  
    for pdf in pdf_list:  
        merger.append(pdf)  
    merger.write(output_file)  
    merger.close()  
    return output_file  
  
# Ensure session state is initialized  
if 'conflict_results' not in st.session_state:  
    st.session_state.conflict_results = None  
if 'foundational_claim' not in st.session_state:  
    st.session_state.foundational_claim = None  
if 'figure_analysis' not in st.session_state:  
    st.session_state.figure_analysis = None  
if 'filed_application_analysis' not in st.session_state:  
    st.session_state.filed_application_analysis = None  
if 'cited_documents' not in st.session_state:  
    st.session_state.cited_documents = None  
if 'pending_claims_analysis' not in st.session_state:  
    st.session_state.pending_claims_analysis = None  
if 'pending_claims_available' not in st.session_state:  
    st.session_state.pending_claims_available = "No"  
if 'domain' not in st.session_state:  
    st.session_state.domain = None  
if 'expertise' not in st.session_state:  
    st.session_state.expertise = None  
if 'style' not in st.session_state:  
    st.session_state.style = None  
if 'filed_application_name' not in st.session_state:  
    st.session_state.filed_application_name = None  
  
# Display the logo and title  
st.image("AFS Innovation Logo.png", width=200)  
st.title("Patent Analyzer")  
  
# Step 1: Upload Examiner Document and Check Conflicts  
with st.expander("Step 1: Office Action", expanded=True):  
    st.write("### Upload the Examiner Document and Check for Conflicts")  
    uploaded_examiner_file = st.file_uploader("Upload Examiner Document", type=["pdf", "docx"])  
    conflicts_clicked = st.button("Check for Conflicts")  
  
    if conflicts_clicked:  
        if uploaded_examiner_file is not None:  
            temp_file_path = "temp_examiner.pdf" if uploaded_examiner_file.type == "application/pdf" else "temp_examiner.docx"  
            with open(temp_file_path, "wb") as f:  
                f.write(uploaded_examiner_file.read())  
  
            if uploaded_examiner_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  
                temp_pdf_path = "temp_examiner_converted.pdf"  
                pdf_path = convert_docx_to_pdf(temp_file_path, temp_pdf_path)  
                if pdf_path:  
                    temp_file_path = pdf_path  
                else:  
                    st.error("Failed to convert DOCX to PDF.")  
  
            if os.path.exists(temp_file_path):  
                extracted_examiner_text = extract_text_from_pdf(temp_file_path)  
                if extracted_examiner_text:  
                    domain, expertise, style = determine_domain_expertise(extracted_examiner_text)  
                    if domain and expertise and style:  
                        st.session_state.domain = domain  
                        st.session_state.expertise = expertise  
                        st.session_state.style = style  
  
                        conflict_results_raw = check_for_conflicts(extracted_examiner_text, domain, expertise, style)  
                        if conflict_results_raw:  
                            st.session_state.conflict_results = conflict_results_raw  
                            st.session_state.foundational_claim = conflict_results_raw.get("foundational_claim", "")  
                            st.session_state.cited_documents = conflict_results_raw.get("documents_referenced", [])  
                            st.success("Conflicts checked successfully!")  
                        else:  
                            st.error("Failed to check for conflicts.")  
                    else:  
                        st.error("Failed to determine domain expertise.")  
                else:  
                    st.error("Failed to extract text from the examiner document.")  
                os.remove(temp_file_path)  
                if uploaded_examiner_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  
                    if os.path.exists(temp_pdf_path):  
                        os.remove(temp_pdf_path)  
            else:  
                st.error("Failed to process the uploaded file.")  
        else:  
            st.warning("Please upload the examiner document first.")  
  
# Display Cited Documents Referenced after Step 1  
if st.session_state.get("cited_documents") is not None:  
    st.write("### Cited Documents Referenced:")  
    cited_docs_df = pd.DataFrame(st.session_state.cited_documents, columns=["Document Name"])  
    st.table(cited_docs_df)  
  
# Step 2: Upload Referenced Document and Analyze Figures  
if st.session_state.get("conflict_results") is not None:  
    with st.expander("Step 2: Referenced Documents", expanded=True):  
        st.write("### Upload the Referenced Documents and Analyze Figures")  
        uploaded_ref_files = st.file_uploader("", type="pdf", key="referenced", accept_multiple_files=True)  
        analyze_figures_clicked = st.button("Analyze Figures and Cited Text")  
  
        if analyze_figures_clicked:  
            if uploaded_ref_files:  
                ref_texts = []  
                for uploaded_ref_file in uploaded_ref_files:  
                    with open(f"temp_{uploaded_ref_file.name}", "wb") as f:  
                        f.write(uploaded_ref_file.read())  
                    extracted_ref_text = extract_text_from_pdf(f"temp_{uploaded_ref_file.name}")  
                    ref_texts.append(extracted_ref_text)  
                    os.remove(f"temp_{uploaded_ref_file.name}")  
  
                figure_analysis_results = extract_figures_and_text(  
                    st.session_state.conflict_results, ref_texts,  
                    st.session_state.domain, st.session_state.expertise, st.session_state.style  
                )  
  
                if figure_analysis_results:  
                    st.session_state.figure_analysis = figure_analysis_results  
                    st.success("Figure analysis completed successfully!")  
                else:  
                    st.error("Failed to analyze figures and cited text.")  
            else:  
                st.warning("Please upload the referenced documents first.")  
  
# Step 3: Ask if the Application is Published  
if st.session_state.get("figure_analysis") is not None:  
    with st.expander("Step 3: Application as Filed", expanded=True):  
        st.write("### Is the Application Published?")  
        is_published = st.radio("Select an option:", ("Yes", "No"))  
  
        if is_published == "No":  
            st.write("### Upload the DOCX and PDF to Combine and Analyze")  
            word_file = st.file_uploader("Upload Word document", type=["docx"])  
            pdf_file = st.file_uploader("Upload PDF document", type=["pdf"])  
            combine_and_proceed_clicked = st.button("Combine and Proceed")  
  
            if combine_and_proceed_clicked:  
                if word_file and pdf_file:  
                    with tempfile.TemporaryDirectory() as tmpdirname:  
                        word_path = os.path.join(tmpdirname, word_file.name)  
                        pdf_path = os.path.join(tmpdirname, pdf_file.name)  
  
                        with open(word_path, "wb") as f:  
                            f.write(word_file.getbuffer())  
                        with open(pdf_path, "wb") as f:  
                            f.write(pdf_file.getbuffer())  
  
                        output_pdf_file = os.path.join(tmpdirname, "combined_document.pdf")  
  
                        with st.spinner("Converting Word to PDF..."):  
                            converted_pdf = convert_word_to_pdf(word_path, os.path.join(tmpdirname, "converted.pdf"))  
  
                        if converted_pdf:  
                            with st.spinner("Merging PDFs..."):  
                                merged_pdf = merge_pdfs([converted_pdf, pdf_path], output_pdf_file)  
  
                            st.success("DOCX and PDF have been successfully combined!")  
                            with open(output_pdf_file, "rb") as f:  
                                st.download_button(  
                                    label="Download Combined PDF",  
                                    data=f,  
                                    file_name="combined_document.pdf",  
                                    mime="application/pdf"  
                                )  
  
                            # Proceed with Step 3 as the combined PDF is ready  
                            with open("temp_filed.pdf", "wb") as f:  
                                f.write(f.read())  
                            extracted_filed_app_text = extract_text_from_pdf("temp_filed.pdf")  
                            os.remove("temp_filed.pdf")  
  
                            if extracted_filed_app_text:  
                                st.session_state.filed_application_name = "Published App US20240090598A1.pdf"  
                                filed_app_details = extract_details_from_filed_application(  
                                    extracted_filed_app_text,  
                                    st.session_state.foundational_claim,  
                                    st.session_state.domain,  
                                    st.session_state.expertise,  
                                    st.session_state.style  
                                )  
                                if filed_app_details:  
                                    filed_app_details_json = json.dumps(filed_app_details, indent=2)  
                                    st.session_state.filed_application_analysis = filed_app_details_json  
  
                                    analysis_results = analyze_filed_application(  
                                        filed_app_details_json,  
                                        st.session_state.foundational_claim,  
                                        st.session_state.figure_analysis,  
                                        st.session_state.domain,  
                                        st.session_state.expertise,  
                                        st.session_state.style  
                                    )  
                                    if analysis_results:  
                                        st.session_state.filed_application_analysis = analysis_results  
                                        st.success("Filed application analysis completed successfully!")  
                                        docx_buffer = save_analysis_to_word(analysis_results)  
                                        if docx_buffer:  
                                            filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                            st.download_button(  
                                                label="Download Analysis Results",  
                                                data=docx_buffer,  
                                                file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                                key="filed_application_download"  
                                            )  
                                    else:  
                                        st.error("Failed to analyze the filed application.")  
                                else:  
                                    st.error("Failed to analyze the filed application.")  
                            else:  
                                st.error("Failed to extract text from the filed application document.")  
                        else:  
                            st.error("Failed to convert Word to PDF.")  
                else:  
                    st.warning("Please upload both the DOCX and PDF files.")  
  
        elif is_published == "Yes":  
            uploaded_filed_app = st.file_uploader("Upload Filed Application", type=["pdf"])  
            analyze_filed_app_clicked = st.button("Analyze Filed Application")  
  
            if analyze_filed_app_clicked:  
                if uploaded_filed_app is not None:  
                    with open("temp_filed.pdf", "wb") as f:  
                        f.write(uploaded_filed_app.read())  
                    extracted_filed_app_text = extract_text_from_pdf("temp_filed.pdf")  
                    os.remove("temp_filed.pdf")  
  
                    if extracted_filed_app_text:  
                        st.session_state.filed_application_name = "Published App US20240090598A1.pdf"  
                        filed_app_details = extract_details_from_filed_application(  
                            extracted_filed_app_text,  
                            st.session_state.foundational_claim,  
                            st.session_state.domain,  
                            st.session_state.expertise,  
                            st.session_state.style  
                        )  
                        if filed_app_details:  
                            filed_app_details_json = json.dumps(filed_app_details, indent=2)  
                            st.session_state.filed_application_analysis = filed_app_details_json  
  
                            analysis_results = analyze_filed_application(  
                                filed_app_details_json,  
                                st.session_state.foundational_claim,  
                                st.session_state.figure_analysis,  
                                st.session_state.domain,  
                                st.session_state.expertise,  
                                st.session_state.style  
                            )  
                            if analysis_results:  
                                st.session_state.filed_application_analysis = analysis_results  
                                st.success("Filed application analysis completed successfully!")  
                                docx_buffer = save_analysis_to_word(analysis_results)  
                                if docx_buffer:  
                                    filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                    st.download_button(  
                                        label="Download Analysis Results",  
                                        data=docx_buffer,  
                                        file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                        key="filed_application_download"  
                                    )  
                            else:  
                                st.error("Failed to analyze the filed application.")  
                        else:  
                            st.error("Failed to analyze the filed application.")  
                    else:  
                        st.error("Failed to extract text from the filed application document.")  
                else:  
                    st.warning("Please upload the filed application first.")  
  
# Step 4: Pending Claims  
if st.session_state.get("filed_application_analysis") is not None:  
    with st.expander("Step 4: Pending Claims", expanded=True):  
        st.write("### Do you have a Pending Claims Document to Analyze?")  
        st.session_state.pending_claims_available = st.radio(  
            "Select an option:",  
            ("Yes", "No"),  
            index=0 if st.session_state.pending_claims_available == "Yes" else 1  
        )  
  
        if st.session_state.pending_claims_available == "Yes":  
            st.write("### Upload the Pending Claims Document and Analyze")  
            uploaded_pending_claims_file, analyze_pending_claims_clicked = st.file_uploader("Upload Pending Claims Document", type=["pdf"]), st.button("Analyze Pending Claims")  
  
            if analyze_pending_claims_clicked:  
                if uploaded_pending_claims_file is not None:  
                    with open("temp_pending_claims.pdf", "wb") as f:  
                        f.write(uploaded_pending_claims_file.read())  
                    extracted_pending_claims_text = extract_text_from_pdf("temp_pending_claims.pdf")  
                    os.remove("temp_pending_claims.pdf")  
  
                    if extracted_pending_claims_text:  
                        modified_filed_application_results = extract_and_modify_filed_application(  
                            st.session_state.filed_application_analysis,  
                            extracted_pending_claims_text,  
                            st.session_state.domain,  
                            st.session_state.expertise,  
                            st.session_state.style  
                        )  
                        if modified_filed_application_results:  
                            st.session_state.modified_filed_application_results = modified_filed_application_results  
                            st.success("Modified filed application analysis completed successfully!")  
  
                            pending_claims_analysis_results = analyze_modified_application(  
                                extracted_pending_claims_text,  
                                st.session_state.foundational_claim,  
                                st.session_state.figure_analysis,  
                                modified_filed_application_results,  
                                st.session_state.domain,  
                                st.session_state.expertise,  
                                st.session_state.style  
                            )  
                            if pending_claims_analysis_results:  
                                st.session_state.pending_claims_analysis = pending_claims_analysis_results  
                                st.success("Pending claims analysis completed successfully!")  
  
                                docx_buffer = save_analysis_to_word(pending_claims_analysis_results)  
                                if docx_buffer:  
                                    filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                    st.download_button(  
                                        label="Download Analysis Results",  
                                        data=docx_buffer,  
                                        file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                        key="pending_claims_download"  
                                    )  
                            else:  
                                st.error("Failed to analyze the pending claims.")  
                        else:  
                            st.error("Failed to modify the filed application based on pending claims.")  
                    else:  
                        st.error("Failed to extract text from the pending claims document.")  
                else:  
                    st.warning("Please upload the pending claims document first.")  
  
# Option to download results if there are no pending claims  
if st.session_state.get("filed_application_analysis") and st.session_state.pending_claims_analysis is None:  
    docx_buffer = save_analysis_to_word(st.session_state.filed_application_analysis)  
    if docx_buffer:  
        filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
        st.download_button(  
            label="Download Analysis Results",  
            data=docx_buffer,  
            file_name=f"{filed_application_name}_ANALYSIS.docx",  
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
            key="filed_application_final_download"  
        )  
